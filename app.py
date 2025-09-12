import os
import io
import tempfile
import streamlit as st
import docx
import re
import json
from groq import Groq

# ─────────────────────────────────────────────────────────────────────────────
# 🎨 Streamlit Config & Styles
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="DOCX Generator", layout="wide", initial_sidebar_state="expanded")
st.markdown(
    """
    <style>
    /* Knoppen */
    .stButton>button, .stDownloadButton>button {font-size:18px; font-weight:bold; padding:0.6em 1.2em;}
    /* Headers */
    .big-header {font-size:2.5rem; font-weight:bold; margin-bottom:0.3em;}
    .section-header {font-size:1.75rem; font-weight:600; margin-top:1em; margin-bottom:0.5em;}
    /* Text Areas */
    .stTextArea textarea {font-family:monospace;}
    </style>
    """,
    unsafe_allow_html=True
)

# ─────────────────────────────────────────────────────────────────────────────
# 🔑 Groq Client Initialization
# ─────────────────────────────────────────────────────────────────────────────
def get_groq_client():
    api_key = st.secrets.get("groq", {}).get("api_key", "").strip()
    if not api_key:
        st.sidebar.error("❌ Voeg Groq API key toe in `.streamlit/secrets.toml` onder [groq]")
        st.stop()
    try:
        return Groq(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"❌ Fout bij verbinden met Groq API: {e}")
        st.stop()

groq_client = get_groq_client()

# ─────────────────────────────────────────────────────────────────────────────
# 🔧 Helper Functions
# ─────────────────────────────────────────────────────────────────────────────
def read_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def get_replacements(template_text: str, context_text: str) -> list[dict]:
    prompt = (
        "Gegeven TEMPLATE en CONTEXT, lever JSON-array met objecten {find, replace}."
        f"\n\nTEMPLATE:\n{template_text}\n\nCONTEXT:\n{context_text}"
    )
    resp = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant", temperature=0.2,
        messages=[
            {"role":"system","content":"Antwoord alleen JSON-array, geen extra tekst."},
            {"role":"user","content":prompt}
        ]
    )
    content = resp.choices[0].message.content
    cleaned = re.sub(r"\d+\s*:\s*{", "{", content)
    start, end = cleaned.find('['), cleaned.rfind(']')+1
    json_str = cleaned[start:end] if start!=-1 and end!=-1 else cleaned
    try:
        repls = json.loads(json_str)
    except json.JSONDecodeError:
        repls = []
        lines = cleaned.splitlines()
        for i, ln in enumerate(lines):
            if '"find"' in ln:
                fm = re.search(r'"find"\s*:\s*"([^"]*)"', ln)
                rm = None
                if fm:
                    for nxt in lines[i+1:]:
                        m = re.search(r'"replace"\s*:\s*"([^"]*)"', nxt)
                        if m:
                            rm = m.group(1)
                            break
                if fm and rm:
                    repls.append({"find":fm.group(1), "replace":rm})
    return [r for r in repls if r['find'] and r['find']!=r['replace']]


def apply_replacements(doc_path: str, replacements: list[dict]) -> bytes:
    doc = docx.Document(doc_path)
    def repl(runs):
        if not runs: return
        txt = ''.join(r.text for r in runs)
        for rp in replacements:
            txt = txt.replace(rp['find'], rp['replace'])
        runs[0].text = txt
        for r in runs[1:]: r.text = ''
    for p in doc.paragraphs: repl(p.runs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl(p.runs)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# 🧭 Page Navigation
# ─────────────────────────────────────────────────────────────────────────────
page = st.sidebar.selectbox("🔖 Navigatie", ["Home", "Generator", "Info"])

if page == "Home":
    st.markdown("<div class='big-header'>🏠 Welkom bij de DOCX Generator</div>", unsafe_allow_html=True)
    st.markdown(
        """
        Gebruik deze tool om snel **Word-templates** bij te werken met **nieuwe context**.
        
        - Ga naar **Generator**
        - Upload je **template** en **context**
        - Klik op **Genereer aangepast document**
        - Download en behoud je opmaak!
        """,
        unsafe_allow_html=True
    )

elif page == "Generator":
    st.markdown("<div class='big-header'>🚀 Generator</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<div class='section-header'>📄 Template Upload</div>", unsafe_allow_html=True)
        tpl_file = st.file_uploader("Kies .docx template", type="docx", key="tpl")
        if tpl_file:
            tpl_path = os.path.join(tempfile.mkdtemp(), "template.docx")
            with open(tpl_path, "wb") as f: f.write(tpl_file.getbuffer())
            tpl_text = read_docx(tpl_path)
            st.text_area("Template-inhoud", tpl_text, height=250, key="tpl_pre")
    with col2:
        st.markdown("<div class='section-header'>📝 Context Upload</div>", unsafe_allow_html=True)
        ctx_file = st.file_uploader("Kies .docx/.txt context", type=["docx","txt"], key="ctx")
        if ctx_file:
            tmp_c = tempfile.mkdtemp()
            if ctx_file.type.endswith('document'):
                cpath = os.path.join(tmp_c, "context.docx")
                with open(cpath,"wb") as f: f.write(ctx_file.getbuffer())
                context = read_docx(cpath)
            else:
                context = ctx_file.read().decode('utf-8', errors='ignore')
            st.text_area("Context-inhoud", context, height=250, key="ctx_pre")
    if tpl_file and ctx_file:
        st.markdown("---")
        if st.button("🎉 Genereer aangepast document"):
            tpl_text = read_docx(tpl_path)
            replacements = get_replacements(tpl_text, context)
            st.markdown("<div class='section-header'>✨ Aangepaste onderdelen</div>", unsafe_allow_html=True)
            for rp in replacements:
                st.write(f"• **{rp['find']}** → **{rp['replace']}**")
            out = apply_replacements(tpl_path, replacements)
            st.download_button(
                "⬇️ Download aangepast document",
                data=out,
                file_name="aangepast_template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Upload eerst template en context om te starten.")

elif page == "Info":
    st.markdown("<div class='big-header'>ℹ️ Info & Tips</div>", unsafe_allow_html=True)
    st.markdown(
        """
        **Tips voor optimaal gebruik:**
        - Zorg voor unieke, duidelijke tekstfragmenten.
        - Houd context-bestanden kort en concreet.
        - Controleer altijd de uiteindelijke output.
        - Voor complexe documenten kun je secties apart bijwerken.
        """,
        unsafe_allow_html=True
    )